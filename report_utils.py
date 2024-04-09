from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
import pandas as pd
import plotly.express as px
from io import BytesIO

PLACES = [
    'Khoa Nhi',
    'Khoa HS-HP',
    'Khoa VS-HM',
    'Khoa Phụ',
    'Khoa CC-PT',
    'Khoa Sanh',
    'Khoa Dược',
    'Khoa XN',
    'Khoa CĐHA',
    'Khoa KSNK',
    'Khoa Dinh Dưỡng',
    'Phòng KHTH-TC-TV',
    'Phòng HCQT',
]

CATEGORY = {
    'A':'Sự cố xày ra có thể tạo ra lỗi/sai sót',
    'B':'Sự cố xày ra nhưng chưa thực hiện trên NB',
    'C':'Sự cố xày ra nhưng không gây hại',
    'D':'Sự cố xày ra đòi hỏi phải theo dõi',
    'E':'Sự cố xày ra trên NB gây tổn hại sức khỏe tạm thời đòi hỏi phải can thiệp chuyên môn',
    'F':'Sự cố xày ra trên NB ảnh hưởng tới sức khỏe hoặc kéo dài ngày nằm viện',
    'G':'Sự cố xày ra trên NB dẫn đến tàn tật vĩnh viễn',
    'H':'Sự cố xày ra trên NB đòi hòi phải can thiệp để cứu sống NB',
    'I':'Sự cố xày ra trên NB gây tử vong',
}

HEADING_LIST =[
    'Danh sách sự cố y khoa {term_year}',
    'Biểu đồ phát hiện sự cố theo thời gian xuất hiện',
    'Biểu đồ phát hiện sự cố theo địa điểm xuất hiện',
    'Biểu đồ phát hiện sự cố theo mức độ nguy hại cho người bệnh',
    'Biểu đồ phát hiện sự cố theo nhóm sự cố',
]

def generate_report_data(df:pd.DataFrame, report_year:int, from_month:int, to_month:int) -> dict:

    report_data = {}
    

    df['year'] = df['Thời gian'].map(lambda x: x.year)
    df['month'] = df['Thời gian'].map(lambda x: x.month)
    df = df.loc[(df['year'] == report_year) & (df['month'] >= from_month) & (df['month'] <= to_month)]

    total = df.shape[0]
    if total==0:
        raise Exception('No available data')

    #Page 1 data
    df_1 = df.rename(columns={
        'Tên SC': 'Mô tả sự cố',
        'BC Tự nguyện': 'Báo cáo tự nguyện',
        'KP báo cáo': 'Khoa/Phòng báo cáo',
        'Đã xảy ra trên NB': 'Đã xảy ra trên người bệnh',
        })
    df_1.drop(columns=['Nhóm'],inplace=True)
    styler_1 =  df_1.style\
        .format({'Thời gian': lambda x: "{}".format(x.strftime('%d-%m-%Y')),
                    'Báo cáo tự nguyện': lambda x: 'Tự nguyện' if x==1 else 'Bắt buộc',
                    'Đã xảy ra trên người bệnh': lambda x: 'Có' if x==1 else 'Không',
                    })
    
    report_data['styler_1'] = styler_1

    #Page 2 data
    df_2 = df['Thời gian'].map(lambda x: x.month).value_counts().rename_axis('Tháng').reset_index(name='Số sự cố')
    df_2['Tỷ lệ %'] = df_2['Số sự cố'].map(lambda x: int(x)*100/total)

    for month in range(from_month,to_month+1):
        if month not in df_2['Tháng'].to_list():
            new_row = {'Tháng': month, 'Số sự cố': 0, 'Tỷ lệ %': 0}
            df_2 = pd.concat([df_2, pd.DataFrame([new_row])], ignore_index=True)
            
    df_2.sort_values(by=['Tháng'], inplace=True)
    df_2['Tháng'] = df_2['Tháng'].map(lambda x: f'Tháng {x}')

    styler_2 =  df_2.style.format({'Tỷ lệ %': "{:.0f}%",})

    fig_2 = px.bar(df_2, x='Tháng', y='Tỷ lệ %',text='Tỷ lệ %', title="Biểu đồ 1: Tỷ lệ sự cố y khoa theo tháng")
    fig_2.update_traces(texttemplate='%{text:.2s}%',textfont_size=12, textangle=0, textposition="outside", cliponaxis=False, marker_color='cadetblue',)

    image_2 = fig_2.to_image(format='png')

    report_data['df_2'] = df_2
    report_data['styler_2'] = styler_2
    report_data['fig_2'] = fig_2
    report_data['image_2'] = image_2

    # Page 3 data
    df_3 = df['Địa điểm'].value_counts().rename_axis('Đơn vị').reset_index(name='Số sự cố')
    df_3['Tỷ lệ %'] = df_3['Số sự cố'].map(lambda x: int(x)*100/total)

    for place in PLACES:
        if place not in df_3['Đơn vị'].to_list():
            new_row = {'Đơn vị': place, 'Số sự cố': 0, 'Tỷ lệ %': 0}
            df_3 = pd.concat([df_3, pd.DataFrame([new_row])], ignore_index=True)
    
    df_3.sort_values(by=['Đơn vị'], inplace=True)
    styler_3 =  df_3.style.format({'Tỷ lệ %': "{:.0f}%",})


    fig_3 = px.bar(df_3, x='Đơn vị', y='Tỷ lệ %',text='Tỷ lệ %', title="Biểu đồ 2: Tỷ lệ sự cố y khoa theo khoa/phòng")
    fig_3.update_traces(texttemplate='%{text:.2s}%',textfont_size=12, textangle=0, textposition="outside", cliponaxis=False, marker_color='slateblue',)
    image_3 = fig_3.to_image(format='png')

    report_data['df_3'] = df_3
    report_data['styler_3'] = styler_3
    report_data['fig_3'] = fig_3
    report_data['image_3'] = image_3

    # Page 4 data
    df_4 = df['Phân loại'].value_counts().reset_index(name='Số sự cố')
    df_4['Tỷ lệ %'] = df_4['Số sự cố'].map(lambda x: int(x)*100/total)
    
    for cate in CATEGORY.keys():
        if cate not in df_4['Phân loại'].to_list():
            new_row = {'Phân loại': cate, 'Số sự cố': 0, 'Tỷ lệ %': 0}
            df_4 = pd.concat([df_4, pd.DataFrame([new_row])], ignore_index=True)

    df_4['Phân loại theo mức độ nguy hiểm cho người bệnh'] = df_4['Phân loại'].map(lambda x: f'Mức {x}')
    df_4['Mô tả'] = df_4['Phân loại'].map(lambda x: CATEGORY[x])
    df_4 = df_4[['Phân loại','Phân loại theo mức độ nguy hiểm cho người bệnh','Mô tả','Số sự cố','Tỷ lệ %']]

    df_4.sort_values(by=['Phân loại'], inplace=True)
    styler_4 =  df_4.style.format({'Tỷ lệ %': "{:.0f}%",}).hide(['Phân loại'], axis=1)

    fig_4 = px.bar(df_4, x='Phân loại theo mức độ nguy hiểm cho người bệnh', y='Tỷ lệ %',text='Tỷ lệ %', title="Biểu đồ 3: Tỷ lệ mức độ nguy hại cho người bệnh")
    fig_4.update_traces(texttemplate='%{text:.2s}%',textfont_size=12, textangle=0, textposition="outside", cliponaxis=False, marker_color='palevioletred')
    image_4 = fig_4.to_image(format='png')

    report_data['df_4'] = df_4
    report_data['styler_4'] = styler_4
    report_data['fig_4'] = fig_4
    report_data['image_4'] = image_4

    # page 5 data
    df_5 = df['Nhóm'].value_counts().rename_axis('Nhóm sự cố').reset_index(name='Số sự cố')
    df_5['Tỷ lệ %'] = df_5['Số sự cố'].map(lambda x: int(x)*100/total)

    df_5.sort_values(by=['Nhóm sự cố'], inplace=True)
    styler_5 =  df_5.style.format({'Tỷ lệ %': "{:.0f}%",})

    fig_5 = px.bar(df_5, x='Nhóm sự cố', y='Số sự cố',text='Số sự cố', title="Biểu đồ 4: Số sự cố")
    fig_5.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False, marker_color='salmon')
    image_5 = fig_5.to_image(format='png')

    report_data['df_5'] = df_5
    report_data['styler_5'] = styler_5
    report_data['fig_5'] = fig_5
    report_data['image_5'] = image_5

    report_data['page'] = [
        (df_1,None),
        (df_2,image_2),
        (df_3,image_3),
        (df_4,image_4),
        (df_5,image_5),
    ]

    return report_data


def create_docx_report(data:dict):

    doc = Document()

    sections = doc.sections
    for section in sections:
        # change orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    doc.add_heading(f'Phân tích sự cố y khoa dạng biểu đồ {data["term_year"]}'.upper(), 1)

    for page in range(len(data['page'])):
        doc.add_heading(HEADING_LIST[page].format(term_year = data['term_year']), 2)
        df, image = data['page'][page]

        t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'Light List Accent 1')
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break()
        
        if image:
            image_stream = BytesIO(image)
            doc.add_picture(image_stream,width=Mm(get_text_width(doc)))
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()

    doc.add_heading('Nhận xét', 2)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    
    doc.add_heading('Giả pháp cải tiến và kiến nghị', 2)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()

    return doc


def get_text_width(document:Document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin)*0.75/ 36000

def show_named_plotly_colours():
    """
    function to display to user the colours to match plotly's named
    css colours.

    Reference:
        #https://community.plotly.com/t/plotly-colours-list/11730/3

    Returns:
        plotly dataframe with cell colour to match named colour name

    """
    s='''
        aliceblue, antiquewhite, aqua, aquamarine, azure,
        beige, bisque, black, blanchedalmond, blue,
        blueviolet, brown, burlywood, cadetblue,
        chartreuse, chocolate, coral, cornflowerblue,
        cornsilk, crimson, cyan, darkblue, darkcyan,
        darkgoldenrod, darkgray, darkgrey, darkgreen,
        darkkhaki, darkmagenta, darkolivegreen, darkorange,
        darkorchid, darkred, darksalmon, darkseagreen,
        darkslateblue, darkslategray, darkslategrey,
        darkturquoise, darkviolet, deeppink, deepskyblue,
        dimgray, dimgrey, dodgerblue, firebrick,
        floralwhite, forestgreen, fuchsia, gainsboro,
        ghostwhite, gold, goldenrod, gray, grey, green,
        greenyellow, honeydew, hotpink, indianred, indigo,
        ivory, khaki, lavender, lavenderblush, lawngreen,
        lemonchiffon, lightblue, lightcoral, lightcyan,
        lightgoldenrodyellow, lightgray, lightgrey,
        lightgreen, lightpink, lightsalmon, lightseagreen,
        lightskyblue, lightslategray, lightslategrey,
        lightsteelblue, lightyellow, lime, limegreen,
        linen, magenta, maroon, mediumaquamarine,
        mediumblue, mediumorchid, mediumpurple,
        mediumseagreen, mediumslateblue, mediumspringgreen,
        mediumturquoise, mediumvioletred, midnightblue,
        mintcream, mistyrose, moccasin, navajowhite, navy,
        oldlace, olive, olivedrab, orange, orangered,
        orchid, palegoldenrod, palegreen, paleturquoise,
        palevioletred, papayawhip, peachpuff, peru, pink,
        plum, powderblue, purple, red, rosybrown,
        royalblue, saddlebrown, salmon, sandybrown,
        seagreen, seashell, sienna, silver, skyblue,
        slateblue, slategray, slategrey, snow, springgreen,
        steelblue, tan, teal, thistle, tomato, turquoise,
        violet, wheat, white, whitesmoke, yellow,
        yellowgreen
        '''
    li=s.split(',')
    li=[l.replace('\n','') for l in li]
    li=[l.replace(' ','') for l in li]

    import pandas as pd
    import plotly.graph_objects as go

    df=pd.DataFrame.from_dict({'colour': li})
    fig = go.Figure(data=[go.Table(
      header=dict(
        values=["Plotly Named CSS colours"],
        line_color='black', fill_color='white',
        align='center', font=dict(color='black', size=14)
      ),
      cells=dict(
        values=[df.colour],
        line_color=[df.colour], fill_color=[df.colour],
        align='center', font=dict(color='black', size=11)
      ))
    ])

    fig.show()